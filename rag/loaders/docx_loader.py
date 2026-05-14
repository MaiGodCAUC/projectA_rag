"""
DOCX 文档解析器

使用 python-docx 提取段落和表格。
Word 文档中的表格保留为结构化 TableData。

依赖：
    pip install python-docx

----------------------------------------------------------------------
## 你需要自己写的部分

本文件由工程辅助生成，建议你理解后自己手写一遍。

学习重点：
1. python-docx 的基本 API：Document()、paragraphs、tables、rows、cells
2. 如何从 Word 内置标题样式中提取章节结构（Heading 1 → level=1）
3. 和 MDLoader 相同的表格提取思路，但 API 不同 —— 训练技术迁移能力

如果你要手写：
- 核心逻辑直观清晰：读段落 → 读表格 → 组装 ParsedDocument
- _extract_sections_from_docx 是亮点 —— Word 标题样式映射到章节层级
- 注意 python-docx 中 style.name 可能是 None，需要防御式检查

没有 TODO(用户) 标记 —— 本文件逻辑完整。
----------------------------------------------------------------------
"""

# ---------------------------------------------------------------------------
# 导入依赖
# ---------------------------------------------------------------------------

# datetime: 生成解析时间戳
from datetime import datetime

# 导入三个数据模型
# ParsedDocument: 解析结果的统一容器
# SectionMeta: 章节信息
# TableData: 结构化表格
from rag.models import ParsedDocument, SectionMeta, TableData


class DocxLoader:
    """DOCX（Word）解析器

    职责：接收 .docx 文件路径，输出 ParsedDocument 对象。

    和 MDLoader 的设计模式完全一致：
    - 对外只暴露 parse() 一个方法
    - 内部逻辑拆分到 _extract_* 私有方法
    - 返回统一的 ParsedDocument 格式

    区别在于 API：
    - MDLoader 用正则表达式解析纯文本
    - DocxLoader 用 python-docx 库解析结构化 XML（.docx 本质是 ZIP 包）
    """

    def parse(self, file_path: str) -> ParsedDocument:
        """解析 DOCX 文件为 ParsedDocument

        三步走：
        1. 提取纯文本段落
        2. 提取章节结构（Word 标题样式）
        3. 提取表格

        Args:
            file_path: DOCX 文件路径

        Returns:
            ParsedDocument 对象
        """
        # 懒加载：只在解析 DOCX 时才导入 python-docx
        # Document 是 python-docx 的核心类，代表整个 Word 文档
        from docx import Document

        # 打开 Word 文档
        # Document(file_path) 解析 .docx 文件（本质是解压 ZIP + 解析 XML）
        doc = Document(file_path)

        # ---- 步骤 1：提取纯文本段落 ----
        # doc.paragraphs 返回文档中所有段落的列表
        # p.text 获取段落的纯文本内容（去除格式标记）
        # if p.text.strip() 过滤掉空白段落（只含空格/换行的段落）
        # 列表推导式：对每个非空段落取 text 属性
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # 将所有段落用双换行符拼接
        # "\n\n".join() 保留段落间的空行分隔，让后续切片知道这是不同段落
        raw_text = "\n\n".join(paragraphs)

        # ---- 步骤 2：提取章节结构 ----
        # 从 Word 内置标题样式（Heading 1/2/3）中提取章节层级
        sections = self._extract_sections_from_docx(doc)

        # ---- 步骤 3：提取表格 ----
        # 从 Word 文档中提取所有表格
        tables = self._extract_tables_from_docx(doc)

        # 组装 ParsedDocument 对象
        # file_name 和 file_hash 由 loader.py 的 load_document 统一填充
        return ParsedDocument(
            file_name="",           # 由 load_document() 补充
            file_type="docx",       # 固定为 "docx"
            raw_text=raw_text,      # 纯文本段落拼接
            sections=sections,      # Word 标题样式映射的章节结构
            tables=tables,          # 结构化表格列表
            parsed_at=datetime.now().isoformat(),  # 记录解析时间
        )

    # ------------------------------------------------------------------
    # 内部方法
    # ------------------------------------------------------------------

    def _extract_sections_from_docx(self, doc) -> list[SectionMeta]:
        """从 Word 标题样式中提取章节结构

        Word 内置样式 Heading 1/2/3 对应 Markdown 的 #/##/###。
        python-docx 中 style.name 为 'Heading 1' / 'Heading 2' 等。

        ## Word 文档的结构

        .docx 文件本质是一个 ZIP 包，里面是 XML 文件。
        python-docx 解析这些 XML，提供了面向对象的 API：
        - Document → 整个文档
        - Paragraph → 段落，有 style 属性（如 'Heading 1'）
        - Run → 段落内的文本片段（相同格式的连续文字）

        ## 字符位置计算

        我们按顺序遍历所有段落，用 char_pos 累加每个段落的字符数，
        从而近似计算出每个章节标题在全文中的字符位置。
        这里的字符位置是「近似」的，因为：
        - +2 模拟段落间距（实际可能有更多空白）
        - 表格文本不计入段落遍历（表格用单独的 API）

        但这个近似精度对于「章节边界切片」来说足够了。

        Args:
            doc: python-docx 的 Document 对象

        Returns:
            SectionMeta 列表
        """
        # 初始化结果列表
        sections = []

        # char_pos: 当前累积的字符位置计数器
        # 从 0 开始，每遍历一个段落就累加
        char_pos = 0

        # 遍历文档中的所有段落
        for para in doc.paragraphs:
            # 计算当前段落的文本长度
            text_l = len(para.text)

            # 获取段落的样式名称
            # para.style 可能为 None（段落没有设置样式）
            # 用三元表达式防御：para.style.name if para.style else ""
            style = para.style.name if para.style else ""

            # 判断：这个段落是不是标题？
            # style.startswith("Heading ") 匹配 'Heading 1'、'Heading 2' 等
            if style.startswith("Heading "):
                # 从样式名称中提取层级数字
                # style.split() → ['Heading', '1'] 或 ['Heading', '2']
                # [-1] 取最后一个元素 → '1'、'2' 等
                # int() 转为整数
                try:
                    level = int(style.split()[-1])
                except ValueError:
                    # 如果解析失败（理论上不会），默认设为 2（二级标题）
                    level = 2

                # 创建 SectionMeta 对象
                # start_char: 当前累积位置
                # end_char: 当前累积位置 + 标题文本长度（标题本身也算入章节范围）
                sections.append(SectionMeta(
                    title=para.text.strip(),        # 标题文本，去首尾空白
                    level=level,                    # 标题层级（1/2/3）
                    start_char=char_pos,            # 标题起始字符位置
                    end_char=char_pos + text_l,     # 标题结束字符位置
                ))

            # 无论是不是标题，都要累加字符位置
            # text_l + 2：+2 是模拟段落之间的空白间隔（\n\n）
            # 这样下一个段落的 start_char 大约在当前位置 + 当前段落长度 + 2
            char_pos += text_l + 2

        # 返回章节列表
        return sections

    def _extract_tables_from_docx(self, doc) -> list[TableData]:
        """提取 Word 文档中的表格

        python-docx 中的 Table API：
        - doc.tables → 文档中所有表格的列表
        - table.rows → 表格中所有行的列表
        - row.cells → 行中所有单元格的列表
        - cell.text → 单元格的文本内容

        和 PDF/Markdown 解析器一样，提取的表格保留为结构化 TableData，
        不摊平为纯文本。

        Args:
            doc: python-docx 的 Document 对象

        Returns:
            TableData 列表
        """
        # 初始化结果列表
        tables = []

        # 遍历文档中的所有表格
        for table in doc.tables:
            # 过滤：少于 2 行的表格没有数据价值
            # table.rows 是表格的行列表
            if len(table.rows) < 2:
                continue  # 只有 1 行（可能只是表头），跳过

            # ---- 解析表头 ----
            # table.rows[0] 是第一行（表头行）
            # table.rows[0].cells 是这一行的所有单元格
            # 列表推导式：提取每个单元格的文本并去空白
            headers = [cell.text.strip() for cell in table.rows[0].cells]

            # ---- 解析数据行 ----
            # table.rows[1:] 是表头之后的所有行
            # 双层列表推导式：
            #   外层 for row in table.rows[1:] → 遍历每一数据行
            #   内层 for cell in row.cells → 遍历行内的每个单元格
            rows = [
                [cell.text.strip() for cell in row.cells]
                for row in table.rows[1:]
            ]

            # 创建 TableData 对象
            # Word 文档没有页码概念，所以不传 page 参数（默认 None）
            tables.append(TableData(headers=headers, rows=rows))

        # 返回表格列表
        return tables
